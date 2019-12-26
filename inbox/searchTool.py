# 打开文档# coding=utf-8
import codecs

import time

import chardet
from docx import Document
import os, sys
import datetime
import xlrd
import pprint
from selenium import webdriver
from htmlUtils import HtmlUtils
from win32com import client as wc
import traceback
import chardet
from pyCache import SearchCache
import threading
import pythoncom


class search():
    searchCache = None
    file_name_result = []
    file_content_result = {}
    count = 0
    process_list = []

    # 对象初始化掉调用方法
    def __init__(self):
        self.searchCache = SearchCache()
        self.lock = threading.RLock()

    def set_process_list(self, tmp):
        self.process_list = tmp

    def get_process_list(self):
        return self.process_list

    def get_count(self):
        self.lock.acquire()
        try:
            return self.count
        except:
            pass
        finally:
            self.lock.release()

    def set_count(self):
        self.lock.acquire()
        try:
            self.count += 1
        except:
            pass
        finally:
            self.lock.release()

    def get_file_name_result(self):
        self.lock.acquire()
        try:
            return self.file_name_result
        except:
            pass
        finally:
            self.lock.release()

    def set_file_name_result(self, tmp):
        self.lock.acquire()
        try:
            self.file_name_result.append(tmp)
        except:
            pass
        finally:
            self.lock.release()

    def get_file_content_result(self):
        self.lock.acquire()
        try:
            return self.file_content_result
        except:
            pass
        finally:
            self.lock.release()

    def set_file_content_result(self, tmp):
        self.lock.acquire()
        try:
            self.file_content_result.update(tmp)
        except:
            pass
        finally:
            self.lock.release()

    # 搜素docx文档
    def search_docx(self, filename, word):
        i = 0
        tmp_list = []
        try:
            # 打开文档
            #  Document(filename)只支持打开后缀为docx的文档
            document = Document(filename)
            #     print("当前搜素的文件名为：", filename)
            for each in document.paragraphs:
                i += 1
                if each.text.find(word) != -1:
                    #                 print("第", i, "行=========》", each.text)
                    tmp_list.append("第 %s行====》%s" % (i, each.text))
        except:
            pass
        return tmp_list

    # 由于我不需要对文档进行修改，
    def search_doc_quick(self, filename, word):
        tmp_list = []
        # 获取文件的时间戳，判断文件同最近的一次查询是否发生了修改
        # 获取当前文件的时间戳
        current_st_mtime = os.stat(filename).st_mtime
        # 获取最近一次查询的文件的时间错
        # print(self.searchCache.get_cachemap())
        pyCache = self.searchCache.get_pyCache(filename)
        # print(pyCache)
        recently_st_mtime = ""
        recently_cach_file_path = ""
        if pyCache:
            recently_st_mtime = pyCache.split("@")[0]
            recently_cach_file_path = pyCache.split("@")[1]
        # 如果当前文件的时间戳和最近一次查询的时间戳相同，说明文件没有发送变化，直接查询缓存目录下的文件
        # print("current_st_mtime=", current_st_mtime)
        # print("recently_st_mtime", recently_st_mtime)
        # print("recently_cach_file_path=", recently_cach_file_path)
        if str(current_st_mtime) == str(recently_st_mtime):  # 此处注意，要将分割出来的时间戳改为字符串，不然等号不成立
            # print("----------------------")
            tmp_list = self.search_docx(recently_cach_file_path, word)
        else:
            tmp_list = self.search_doc(filename, word)
        return tmp_list

    # 向缓存文件里写入信息
    def update_cache_map(self, filename, docx_filename):
        # print(filename)
        self.searchCache.get_cachemap().update({filename: str(os.stat(filename).st_mtime) + "@" + docx_filename})
        # print("SearchCache.get_cachemap()=", self.searchCache.get_cachemap())

    # 搜素doc文档 该方法相当慢，是把doc文档转为docx温文档后再查询
    def search_doc(self, filename, word):
        tmp_list = []
        docx_filename = ""
        doc = None
        pythoncom.CoInitialize()
        wordApplication = wc.DispatchEx('Word.Application')
        try:
            # 将doc文件转为docx文件处理
            docx_filename = filename.replace("doc", "docx")
            # print("docx_filename0=",docx_filename)
            # 这个地方有点坑，若转为docx后的文件存放路径不和源文件路径一直，这个地方会非常慢，但是我这边为了多次查询效率更高
            # 还是把转为docx的文件缓存到另外一个路径，这样仅仅是第一次查询慢，后面的由于不必再次转换，将会非常快
            docx_filename = self.searchCache.get_cachePath() + "/" + ((os.path.split(docx_filename)[0]).replace("\\",
                                                                                                                "_")).replace(
                "/", "_").replace(
                ":", "_") + "_" + os.path.split(docx_filename)[1]  # 拼接缓存路径
            # print(docx_filename)
            # print("self.searchCache.get_cachePath()",self.searchCache.get_cachePath())
            # print("docx_filename1=",docx_filename)
            # 若是多线程，此处必须加上这行代码，否则会报 pywintypes.com_error: (-2147221008, '尚未调用 CoInitialize。', None, None)的错误
            doc = wordApplication.Documents.Open(filename)  # 生成的docx文件放到缓存路径下
            doc.SaveAs(docx_filename, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
            tmp_list = self.search_docx(docx_filename, word)
            #             # 打开文档
            #             #  Document(filename)只支持打开后缀为docx的文档
            #             document = Document(docx_filename)
            #             #     print("当前搜素的文件名为：", filename)
            #             for each in document.paragraphs:
            #                 i += 1
            #                 if each.text.find(word) != -1:
            #                     #                 print("第", i, "行=========》", each.text)
            #                     tmp_list.append("第 %s行====》%s" % (i, each.text))
            #             # 搜素完了将转换的文件删除
            #             if os.path.exists(docx_filename):
            #                 os.remove(docx_filename)
            # wordApplication.Quit()
        except:
            pass
            traceback.print_exc()
        finally:
            # 将缓存信息写入缓存文件
            pythoncom.CoInitialize()
            if doc:
                try:
                    doc.Close()
                    if wordApplication:
                        wordApplication.Quit()
                except:
                    traceback.print_exc()
            self.update_cache_map(filename, docx_filename)
        return tmp_list

    # 搜素excle文档
    def search_xls(self, filename, word):
        tmp_list = []
        try:
            # 打开excel类型文件
            xlsBook = xlrd.open_workbook(filename)
            # 获取sheet个数
            sheetsNum = xlsBook.nsheets
            # 得到sheet集合
            sheets = xlsBook.sheets()
            #         print("获取的所有sheet为", sheets)
            # 遍历sheet页
            # sheet页是从0开始
            for m in range(0, sheetsNum):
                sheet = sheets[m]
                sheetName = xlsBook.sheet_names()[m]
                #             print("获取的sheet是：",sheet)
                # 获取当前sheet的行数
                rows = sheet.nrows
                # 遍历sheet的每行数据
                # 行数是从0开始
                for n in range(0, rows):
                    row_data = sheet.row_values(n)
                    for each in row_data:
                        #                     print(each)
                        # 由于excel 中有些事数字，这里把所有类型的数据转为字符串，避免异常
                        each = str(each)
                        if each.find(word) != -1:
                            tmp_list.append("%s 中第 %s 行包含 %s" % (sheetName, n, word))
        except:
            pass
        return tmp_list

    def deal_asin_file(self, filename):
        f = open(filename, 'rb+')  # 按字节码流打开文件
        data = f.read()  # 按字节码读取文件
        f.close()
        data_chardet = chardet.detect(data)  # 获取文件的编码信息
        data_encoding = data_chardet["encoding"]  # 获取文件的编码 # 此处注意，若文件的编码为asin，这个方法获取不到文件的编码集，这个地方有点坑
        print(data_encoding)
        data_str = data.decode(data_encoding)
        print(data_str)

    # 搜素txt，java，md，.bat等非二级制文件
    def search_txt(self, filename, word):
        tmp_list = []
        try:
            data = ""
            # 判断并获取txt类型文件编码
            try:
                f = open(filename, 'rb+')
                content = f.read()  # 读取文件内容，content为bytes类型，而非string类型
                f.close()
                content.decode('utf-8').encode('utf-8')
                source_encoding = 'utf-8'
            except:
                try:
                    content.decode('gbk').encode('utf-8')
                    source_encoding = 'gbk'
                except:
                    try:
                        content.decode('gb2312').encode('utf-8')
                        source_encoding = 'gb2312'
                    except:
                        try:
                            content.decode('gb18030').encode('utf-8')
                            source_encoding = 'gb18030'
                        except:
                            try:
                                content.decode('big5').encode('utf-8')
                                source_encoding = 'gb18030'
                            except:
                                content.decode('cp936').encode('utf-8')
                                source_encoding = 'cp936'
            # 获取文件字符串
            # 方法一：直接通过以上得到的文件内容的字节内容content和文件的编码，获取正确的字符串
            data = content.decode(source_encoding)
            if data.find(word) != -1:
                tmp_list.append("%s 中包含 %s" % (filename, word))
            # 方法二：通过上面获取的正确的文件编码集重新打开字符串文件流
            # with open(filename, "r", encoding=source_encoding) as f:
            #     data = f.read()
            #     f.close()
            #     if data.find(word) != -1:
            #         tmp_list.append("%s 中包含 %s" % (filename, word))
        except:
            # traceback.print_exc()
            pass
        return tmp_list

    def search_word(self, filename, word):
        '''
        from docx import Document
        import xlrd
        from win32com import client as wc
        :param filename:
        :param word:
        :return:
        docx文档使用docx api处理
        doc文档使用win32com api处理
        excel使用xlrd api处理
        txt使用open处理
        三种类型的文档处理速度：docx > excel > doc> txt
        '''
        tmp_list = []
        if os.path.splitext(filename)[1] in [".docx"]:
            tmp_list = self.search_docx(filename, word)
        elif os.path.splitext(filename)[1] in [".doc"]:
            tmp_list = self.search_doc_quick(filename, word)
        elif os.path.splitext(filename)[1] in [".xlsx", ".xls"]:
            tmp_list = self.search_xls(filename, word)
        # else:
        #     tmp_list = self.search_txt(filename, word)
        result = {}
        if tmp_list != []:
            result = {filename: tmp_list}
        return result

    def get_process_files(self, root_dir):
        """process all files in directory"""
        process_list = []
        try:
            cur_dir = os.path.abspath(root_dir)
            file_list = os.listdir(cur_dir)

            for file in file_list:
                fullfile = cur_dir + "\\" + file
                if os.path.isfile(fullfile):
                    process_list.append(fullfile)
                elif os.path.isdir(fullfile):
                    dir_extra_list = self.get_process_files(fullfile)
                    if len(dir_extra_list) != 0:
                        for x in dir_extra_list:
                            process_list.append(x)
        except:
            pass
        return process_list

    # 查找文件名中文保护指定字符串
    def find_files_byname(self, process_list, kword):
        result_list = []
        try:
            for each in process_list:
                if os.path.split(each)[1].find(kword) != -1:
                    result_list.append(each.replace("\\\\", "\\"))
        except:
            traceback.print_exc()
        return result_list

    def find_files_content_thread(self, process_list, kword, len_process_list):
        for file in process_list:
            self.set_count()
            i = self.get_count()
            a = i / len_process_list
            if i > 10000 and i % 10000 == 0:
                print("总文件个数为：", len_process_list, "  当前正在搜索第", i, "个文件")
                print("搜素进度为：", "%.2f%%" % (a * 100))
            elif i > 1000 and i % 1000 == 0:
                print("总文件个数为：", len_process_list, "  当前正在搜索第", i, "个文件")
                print("搜素进度为：", "%.2f%%" % (a * 100))
            elif i <= 100:
                print("总文件个数为：", len_process_list, "  当前正在搜索第", i, "个文件")
                print("搜素进度为：", "%.2f%%" % (a * 100))
            if os.path.split(file)[1].find(kword) != -1:
                self.set_file_name_result(file)
            # 判断文件中的文件内容是否可搜索
            if self.can_be_search(file):
                self.set_file_content_result(self.search_word(file, kword))  # 更新文件内容结果集

    # 　多线程查询文件内容
    def find_files_content_quick(self, process_list, kword):
        len_process_list = len(process_list)
        tmp_list_size = len_process_list // 4

        t1 = threading.Thread(target=self.find_files_content_thread,
                              args=(process_list[0: tmp_list_size], kword, len_process_list,))
        t2 = threading.Thread(target=self.find_files_content_thread,
                              args=(process_list[tmp_list_size: 2 * tmp_list_size], kword, len_process_list,))
        t3 = threading.Thread(target=self.find_files_content_thread,
                              args=(
                                  process_list[2 * tmp_list_size: 3 * tmp_list_size], kword,
                                  len_process_list,))
        t4 = threading.Thread(target=self.find_files_content_thread,
                              args=(process_list[3 * tmp_list_size:], kword, len_process_list,))
        t1.start()
        t2.start()
        t3.start()
        t4.start()
        t1.join()
        t2.join()
        t3.join()
        t4.join()
        return [self.get_file_name_result(), self.get_file_content_result()]

    def find_files_content(self, process_list, kword):
        # pprint.pprint("所有文件为%s" % process_list)
        i = 0
        # 获取可查询文件内容的文件列表
        # canBeSerach_file_list = []
        # canBeSerach_file_list = getSearch_file_list(process_list)
        # # pprint.pprint("可查询的文件为%s" % canBeSerach_file_list)
        # 文件内容结果集
        file_content_result = {}
        # 　文件名结果集
        tmp_list = []
        for file in process_list:
            i += 1
            a = i / len(process_list)
            if i > 10000 and i % 10000 == 0:
                print("总文件个数为：", len(process_list), "  当前正在搜索第", i, "个文件")
                print("搜素进度为：", "%.2f%%" % (a * 100))
            elif i > 1000 and i % 1000 == 0:
                print("总文件个数为：", len(process_list), "  当前正在搜索第", i, "个文件")
                print("搜素进度为：", "%.2f%%" % (a * 100))
            elif i <= 100:
                print("总文件个数为：", len(process_list), "  当前正在搜索第", i, "个文件")
                print("搜素进度为：", "%.2f%%" % (a * 100))
            if os.path.split(file)[1].find(kword) != -1:
                tmp_list.append(file)
            # 判断文件中的文件内容是否可搜索
            if self.can_be_search(file):
                file_content_result.update(self.search_word(file, kword))  # 更新文件内容结果集
        file_name_result = tmp_list  # 更新文件名结果集
        return [file_name_result, file_content_result]

    def find_files(self, root_dir, kword, output_path, only_file_name=False):
        start_time = datetime.datetime.now()
        pprint.pprint("开始搜素，请耐心等待..............")
        # 查找传入目录下的所有文件
        process_list = self.get_process_files(root_dir)
        file_name_result = {}
        file_content_result = {}
        if only_file_name:
            file_name_result = self.find_files_byname(process_list, kword)
        else:
            # file_result = self.find_files_content(process_list, kword)
            file_result = self.find_files_content_quick(process_list, kword)
            file_name_result = file_result[0]
            file_content_result = file_result[1]

        end_time = datetime.datetime.now()
        # 搜素耗时
        cost_time = (end_time - start_time).seconds
        cost_time_str = str(cost_time) + "s"
        # pprint.pprint("耗时 %s秒" % cost_time)
        # pprint.pprint("以下文件中包含 %s字符%s" % (word, result))
        self.deal_with_result(root_dir, word, output_path, cost_time_str, file_name_result, file_content_result)

    # 判断文件是否可搜素
    def can_be_search(self, file):
        if os.path.splitext(file)[1] in [".doc", ".docx", ".xlsx",
                                         ".xls"]:  # 二进制文件中"docx", "xlsx", "xls","doc"后缀的文件也可查
            return True
        elif not self.is_binary_file(file):
            return True

    # 获取可查询文件内容的文件列表，支持查询doc，docx，xls，xlsx和txt，md，java，bat文件
    def getSearch_file_list(self, process_list):
        search_file_list = []
        for tmpFile in process_list:
            # print("~~~~~", tmpFile)
            # 判断每个文件是否为二进制文件
            if not self.is_binary_file(tmpFile):
                search_file_list.append(tmpFile)
            elif os.path.splitext(tmpFile)[1] in [".doc", ".docx", ".xlsx",
                                                  ".xls"]:  # 二进制文件中"docx", "xlsx", "xls","doc"后缀的文件也可查
                search_file_list.append(tmpFile)
        return search_file_list

    # 判断一个文件是否为二进制文件
    def is_binary_file(self, file_path):
        _TEXT_BOMS = (
            codecs.BOM_UTF16_BE,
            codecs.BOM_UTF16_LE,
            codecs.BOM_UTF32_BE,
            codecs.BOM_UTF32_LE,
            codecs.BOM_UTF8,
        )
        try:
            with open(file_path, 'rb') as file:
                initial_bytes = file.read(8192)
                file.close()
                flag = not any(initial_bytes.startswith(bom) for bom in _TEXT_BOMS) and b'\0' in initial_bytes
                return flag
        except:
            # traceback.print_exc()
            pass

    # 处理查询出的结果集 result类型为{[],str:[]}
    def deal_with_result(self, root_dir, word, output_path, cost_time, file_name_result=[], file_content_result={}):
        dect_keys = file_content_result.keys()
        file_name_list = ""
        file_content_list = ""
        # 　文件名搜素结果集
        for each in file_name_result:
            file_name_list += "<a href='file:///%s'>%s</a></br>" % (each, each)
        # 文件内容搜索結果集
        for each_1 in dect_keys:  # 遍历文件名
            file_content_list += "<a href='file:///%s'>%s</a></br>" % (each_1, each_1)
            for each_2 in file_content_result.get(each_1):  # 遍历每个文件下匹配到的内容的大致位置
                file_content_list += "<p>%s</p>" % each_2
        message = """<html>
        <meta http-equiv="Content-Type" content="text/html; charset=UTF-8" />
        <style type='text/css'>
        p{text-indent:2em}
        </style>
        <head></head>
        <body>
        <span>在 %s 文件夹下搜素 %s 的结果集如下，耗费时长为：%s</span></br>
        <span>文件名包含 '%s' 的有 %s 个:</span></br>
        %s
        <span>文件内容包含 '%s' 的有  %s 个:</span><br>
        %s
        </body>
        </html>""" % (
            root_dir, word, cost_time, word, len(file_name_result), file_name_list, word, len(file_content_result),
            file_content_list)

        file_path = output_path
        HtmlUtils().create_html(message, file_path)  # 生成的文件请使用火狐浏览器打开，因为火狐浏览器能以上写法火狐浏览器能调用资源管理器打开文件
        HtmlUtils().open_browser_by_webbrowser(file_path)


if __name__ == '__main__':
    # 文件根目录
    # root_dir=sys.argv[1]
    root_dir = "E:/java/ideaWorkspace/pythonTest/tmp"
    # root_dir = "D:/java/eclipse-workspace/python/tmp"
    # 要搜索的关键字
    #     word = sys.argv[2]
    # 工单对应修改模块 Memory
    word = "简书"
    output_path = "E:\\java\\ideaWorkspace\\pythonTest\\tmp\\result.html"
    search().find_files(root_dir, word, output_path, False)
    # E:\\jonesWorkSpace\\笔记\\test2\\new 9.txt
    # E:\\jonesWorkSpace\\笔记\\test2\\常用工作薄_常用网站密码.txt
    # deal_asin_file("E:\\jonesWorkSpace\\笔记\\test2\\new 9.txt")
    # print(search().search_doc_quick("E:/java/ideaWorkspace/pythonTest/tmp/新建 DOC 文档.doc", word));

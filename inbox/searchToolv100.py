# 打开文档# coding=utf-8
import codecs

import time

import chardet
from docx import Document
import os, sys
import datetime
import xlrd
import pprint
from selenium import webdriver
from htmlUtils import HtmlUtils
from win32com import client as wc
import traceback
import chardet


# import htmlUtils


# 查找word,excel,txt等可以被查找的文件中是否包含指定字符串，并列出文件名
# 此处注意，search_word方法中有个带有默认参数的关键字形参，该形参的默认实参属于该函数，
# 所以既该方法如果第一次调用，实参对象就会被创建，同一个主线程多次调用，实参会一直存在，不会被销毁
# 如果实参是一个list，而方法里又使用list.append()方法，那么该对象会越来越大，从不会销毁
# 所以对于此得慎用
def search_word_unRecommend(filename, word, result=[]):
    i = 0
    tmp_list = []
    if os.path.splitext(filename)[1] in [".docx"]:
        try:
            # 打开文档
            #  Document(filename)只支持打开后缀为docx的文档
            document = Document(filename)
            #     print("当前搜素的文件名为：", filename)
            for each in document.paragraphs:
                i += 1
                if each.text.find(word) != -1:
                    #                 print("第", i, "行=========》", each.text)
                    tmp_list.append("第 %s行====》%s" % (i, each.text))
        except:
            pass
    elif os.path.splitext(filename)[1] in [".xlsx", ".xls"]:
        try:
            # 打开excel类型文件
            xlsBook = xlrd.open_workbook(filename)
            # 获取sheet个数
            sheetsNum = xlsBook.nsheets
            # 得到sheet集合
            sheets = xlsBook.sheets()
            #         print("获取的所有sheet为", sheets)
            # 遍历sheet页
            # sheet页是从0开始
            for m in range(0, sheetsNum):
                sheet = sheets[m]
                sheetName = xlsBook.sheet_names()[m]
                #             print("获取的sheet是：",sheet)
                # 获取当前sheet的行数
                rows = sheet.nrows
                # 遍历sheet的每行数据
                # 行数是从0开始
                for n in range(0, rows):
                    row_data = sheet.row_values(n)
                    for each in row_data:
                        #                     print(each)
                        # 由于excel 中有些事数字，这里把所有类型的数据转为字符串，避免异常
                        each = str(each)
                        if word in each:
                            tmp_list.append("%s 中第 %s 行包含 %s" % (sheetName, n, word))
        except:
            traceback.print_exc()
    else:
        try:
            # 判断并获取txt类型文件编码
            try:
                f = open(filename, 'rb+')
                content = f.read()  # 读取文件内容，content为bytes类型，而非string类型
                content.decode('utf-8').encode('utf-8')
                source_encoding = 'utf-8'
            except:
                try:
                    content.decode('gbk').encode('utf-8')
                    source_encoding = 'gbk'
                except:
                    try:
                        content.decode('gb2312').encode('utf-8')
                        source_encoding = 'gb2312'
                    except:
                        try:
                            content.decode('gb18030').encode('utf-8')
                            source_encoding = 'gb18030'
                        except:
                            try:
                                content.decode('big5').encode('utf-8')
                                source_encoding = 'gb18030'
                            except:
                                content.decode('cp936').encode('utf-8')
                                source_encoding = 'cp936'
            # 打开txt类型的文件
            with open(filename, "r", encoding=source_encoding) as f:
                data = f.read()
                #             print(data)
                if word in data:
                    tmp_list.append("%s 中包含 %s" % (filename, word))
        except:
            traceback.print_exc()

    if tmp_list != []:
        result.append({filename: tmp_list})
    return result;


# 搜素docx文档
def search_docx(filename, word):
    i = 0
    tmp_list = []
    try:
        # 打开文档
        #  Document(filename)只支持打开后缀为docx的文档
        document = Document(filename)
        #     print("当前搜素的文件名为：", filename)
        for each in document.paragraphs:
            i += 1
            if each.text.find(word) != -1:
                #                 print("第", i, "行=========》", each.text)
                tmp_list.append("第 %s行====》%s" % (i, each.text))
    except:
        pass
    return tmp_list


# 搜素doc文档 该方法相当慢，是把doc文档转为docx温文档后再查询
def search_doc(filename, word):
    i = 0
    tmp_list = []
    try:
        # 将doc文件转为docx文件处理
        docx_filename = filename.replace("doc", "docx")
        word_tmp = wc.Dispatch('Word.Application')
        doc = word_tmp.Documents.Open(filename)  # 目标路径下的文件
        doc.SaveAs(docx_filename, 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
        doc.Close()
        # 打开文档
        #  Document(filename)只支持打开后缀为docx的文档
        document = Document(docx_filename)
        #     print("当前搜素的文件名为：", filename)
        for each in document.paragraphs:
            i += 1
            if each.text.find(word) != -1:
                #                 print("第", i, "行=========》", each.text)
                tmp_list.append("第 %s行====》%s" % (i, each.text))
        # 搜素完了将转换的文件删除
        if os.path.exists(docx_filename):
            os.remove(docx_filename)
        word_tmp.Quit()
    except:
        pass
    return tmp_list


# 搜素excle文档
def search_xls(filename, word):
    tmp_list = []
    try:
        # 打开excel类型文件
        xlsBook = xlrd.open_workbook(filename)
        # 获取sheet个数
        sheetsNum = xlsBook.nsheets
        # 得到sheet集合
        sheets = xlsBook.sheets()
        #         print("获取的所有sheet为", sheets)
        # 遍历sheet页
        # sheet页是从0开始
        for m in range(0, sheetsNum):
            sheet = sheets[m]
            sheetName = xlsBook.sheet_names()[m]
            #             print("获取的sheet是：",sheet)
            # 获取当前sheet的行数
            rows = sheet.nrows
            # 遍历sheet的每行数据
            # 行数是从0开始
            for n in range(0, rows):
                row_data = sheet.row_values(n)
                for each in row_data:
                    #                     print(each)
                    # 由于excel 中有些事数字，这里把所有类型的数据转为字符串，避免异常
                    each = str(each)
                    if each.find(word) != -1:
                        tmp_list.append("%s 中第 %s 行包含 %s" % (sheetName, n, word))
    except:
        pass
    return tmp_list


def deal_asin_file(filename):
    f = open(filename, 'rb+')  # 按字节码流打开文件
    data = f.read()  # 按字节码读取文件
    f.close()
    data_chardet = chardet.detect(data)  # 获取文件的编码信息
    data_encoding = data_chardet["encoding"]  # 获取文件的编码 # 此处注意，若文件的编码为asin，这个方法获取不到文件的编码集，这个地方有点坑
    print(data_encoding)
    data_str = data.decode(data_encoding)
    print(data_str)


# 搜素txt，java，md，.bat等非二级制文件
def search_txt(filename, word):
    tmp_list = []
    try:
        data = ""
        # 判断并获取txt类型文件编码
        try:
            f = open(filename, 'rb+')
            content = f.read()  # 读取文件内容，content为bytes类型，而非string类型
            f.close()
            content.decode('utf-8').encode('utf-8')
            source_encoding = 'utf-8'
        except:
            try:
                content.decode('gbk').encode('utf-8')
                source_encoding = 'gbk'
            except:
                try:
                    content.decode('gb2312').encode('utf-8')
                    source_encoding = 'gb2312'
                except:
                    try:
                        content.decode('gb18030').encode('utf-8')
                        source_encoding = 'gb18030'
                    except:
                        try:
                            content.decode('big5').encode('utf-8')
                            source_encoding = 'gb18030'
                        except:
                            content.decode('cp936').encode('utf-8')
                            source_encoding = 'cp936'
        # 获取文件字符串
        # 方法一：直接通过以上得到的文件内容的字节内容content和文件的编码，获取正确的字符串
        data = content.decode(source_encoding)
        if data.find(word) != -1:
            tmp_list.append("%s 中包含 %s" % (filename, word))
        # 方法二：通过上面获取的正确的文件编码集重新打开字符串文件流
        # with open(filename, "r", encoding=source_encoding) as f:
        #     data = f.read()
        #     f.close()
        #     if data.find(word) != -1:
        #         tmp_list.append("%s 中包含 %s" % (filename, word))
    except:
        # traceback.print_exc()
        pass
    return tmp_list


def search_word(filename, word):
    '''
    from docx import Document
    import xlrd
    from win32com import client as wc
    :param filename:
    :param word:
    :return:
    docx文档使用docx api处理
    doc文档使用win32com api处理
    excel使用xlrd api处理
    txt使用open处理
    三种类型的文档处理速度：docx > excel > doc> txt
    '''
    tmp_list = []
    if os.path.splitext(filename)[1] in [".docx"]:
        tmp_list = search_docx(filename, word)
    elif os.path.splitext(filename)[1] in [".doc"]:
        tmp_list = search_doc(filename, word)
    elif os.path.splitext(filename)[1] in [".xlsx", ".xls"]:
        tmp_list = search_xls(filename, word)
    else:
        tmp_list = search_txt(filename, word)
    result = {}
    if tmp_list != []:
        result = {filename: tmp_list}
    return result


def get_process_files(root_dir):
    """process all files in directory"""
    cur_dir = os.path.abspath(root_dir)
    file_list = os.listdir(cur_dir)
    process_list = []
    for file in file_list:
        fullfile = cur_dir + "\\" + file
        if os.path.isfile(fullfile):
            process_list.append(fullfile)
        elif os.path.isdir(fullfile):
            dir_extra_list = get_process_files(fullfile)
            if len(dir_extra_list) != 0:
                for x in dir_extra_list:
                    process_list.append(x)
    return process_list


# 查找文件名中文保护指定字符串
def find_files_byname(process_list, kword):
    result_list = []
    try:
        for each in process_list:
            if os.path.split(each)[1].find(kword) != -1:
                result_list.append(each.replace("\\\\", "\\"))
    except:
        traceback.print_exc()
    return result_list


def find_files_content(process_list, kword):
    # pprint.pprint("所有文件为%s" % process_list)
    i = 0
    # 获取可查询文件内容的文件列表
    # canBeSerach_file_list = []
    # canBeSerach_file_list = getSearch_file_list(process_list)
    # # pprint.pprint("可查询的文件为%s" % canBeSerach_file_list)
    # 文件内容结果集
    file_content_result = {}
    # 　文件名结果集
    tmp_list = []
    for file in process_list:
        i += 1
        a = i / len(process_list)
        if i > 10000 and i % 10000 == 0:
            print("总文件个数为：", len(process_list), "  当前正在搜索第", i, "个文件")
            print("搜素进度为：", "%.2f%%" % (a * 100))
        elif i > 1000 and i % 1000 == 0:
            print("总文件个数为：", len(process_list), "  当前正在搜索第", i, "个文件")
            print("搜素进度为：", "%.2f%%" % (a * 100))
        elif i <= 100:
            print("总文件个数为：", len(process_list), "  当前正在搜索第", i, "个文件")
            print("搜素进度为：", "%.2f%%" % (a * 100))
        if file == "E:\\jonesWorkSpace\\tmp\\简书.docx":
            print("os.path.split(file)[1].find(kword)=", os.path.split(file)[1].find(kword))
        if os.path.split(file)[1].find(kword) != -1:
            tmp_list.append(file)
        # 判断文件中的文件内容是否可搜索
        if can_be_search(file):
            file_content_result.update(search_word(file, kword))  # 更新文件内容结果集
    file_name_result = tmp_list  # 更新文件名结果集
    return [file_name_result, file_content_result]


def find_files(root_dir, kword, output_path, only_file_name=False):
    start_time = datetime.datetime.now()
    pprint.pprint("开始搜素，请耐心等待..............")
    # 查找传入目录下的所有文件
    process_list = get_process_files(root_dir)
    file_name_result = {}
    file_content_result = {}
    if only_file_name:
        file_name_result = find_files_byname(process_list, kword)
    else:
        file_result = find_files_content(process_list, kword)
        file_name_result = file_result[0]
        file_content_result = file_result[1]

    end_time = datetime.datetime.now()
    # 搜素耗时
    cost_time = (end_time - start_time).seconds
    cost_time_str = str(cost_time) + "s"
    # pprint.pprint("耗时 %s秒" % cost_time)
    # pprint.pprint("以下文件中包含 %s字符%s" % (word, result))
    deal_with_result(root_dir, word, output_path, cost_time_str, file_name_result, file_content_result)


# 判断文件是否可搜素
def can_be_search(file):
    if os.path.splitext(file)[1] in [".doc", ".docx", ".xlsx",
                                     ".xls"]:  # 二进制文件中"docx", "xlsx", "xls","doc"后缀的文件也可查
        return True
    elif not is_binary_file(file):
        return True


# 获取可查询文件内容的文件列表，支持查询doc，docx，xls，xlsx和txt，md，java，bat文件
def getSearch_file_list(process_list):
    search_file_list = []
    for tmpFile in process_list:
        # print("~~~~~", tmpFile)
        # 判断每个文件是否为二进制文件
        if not is_binary_file(tmpFile):
            search_file_list.append(tmpFile)
        elif os.path.splitext(tmpFile)[1] in [".doc", ".docx", ".xlsx",
                                              ".xls"]:  # 二进制文件中"docx", "xlsx", "xls","doc"后缀的文件也可查
            search_file_list.append(tmpFile)
    return search_file_list


# 判断一个文件是否为二进制文件
def is_binary_file(file_path):
    _TEXT_BOMS = (
        codecs.BOM_UTF16_BE,
        codecs.BOM_UTF16_LE,
        codecs.BOM_UTF32_BE,
        codecs.BOM_UTF32_LE,
        codecs.BOM_UTF8,
    )
    try:
        with open(file_path, 'rb') as file:
            initial_bytes = file.read(8192)
            file.close()
            flag = not any(initial_bytes.startswith(bom) for bom in _TEXT_BOMS) and b'\0' in initial_bytes
            return flag
    except:
        # traceback.print_exc()
        pass


# 处理查询出的结果集 result类型为{str:[],str:[]}
def deal_with_result(root_dir, word, output_path, cost_time, file_name_result=[], file_content_result={}):
    dect_keys = file_content_result.keys()
    file_name_list = ""
    file_content_list = ""
    # 　文件名搜素结果集
    for each in file_name_result:
        file_name_list += "<a href='file:///%s'>%s</a></br>" % (each, each)
    # 文件内容搜索結果集
    for each_1 in dect_keys:  # 遍历文件名
        file_content_list += "<a href='file:///%s'>%s</a></br>" % (each_1, each_1)
        for each_2 in file_content_result.get(each_1):  # 遍历每个文件下匹配到的内容的大致位置
            file_content_list += "<p>%s</p>" % each_2
    message = """<html>
    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8" />
    <style type='text/css'>
    p{text-indent:2em}
    </style>
    <head></head>
    <body>
    <span>在 %s 文件夹下搜素 %s 的结果集如下，耗费时长为：%s</span></br>
    <span>文件名包含 '%s' 的有 %s 个:</span></br>
    %s
    <span>文件内容包含 '%s' 的有  %s 个:</span><br>
    %s
    </body>
    </html>""" % (
        root_dir, word, cost_time, word, len(file_name_result), file_name_list, word, len(file_content_result),
        file_content_list)

    file_path = output_path
    HtmlUtils().create_html(message, file_path)  # 生成的文件请使用火狐浏览器打开，因为火狐浏览器能以上写法火狐浏览器能调用资源管理器打开文件
    HtmlUtils().open_browser_by_webbrowser(file_path)


if __name__ == '__main__':
    # 文件根目录
    # root_dir=sys.argv[1]
    root_dir = "E:/jonesWorkSpace"
    # root_dir = "E:\\jonesWorkSpace"
    # 要搜索的关键字
    #     word = sys.argv[2]
    # 工单对应修改模块 Memory
    word = "简书"
    output_path = "E:\\java\\ideaWorkspace\\pythonTest\\tmp\\result.html"
    find_files(root_dir, word, output_path, False)
    # E:\\jonesWorkSpace\\笔记\\test2\\new 9.txt
    # E:\\jonesWorkSpace\\笔记\\test2\\常用工作薄_常用网站密码.txt
    # deal_asin_file("E:\\jonesWorkSpace\\笔记\\test2\\new 9.txt")
